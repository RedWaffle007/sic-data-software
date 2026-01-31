import pandas as pd
import sys
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_BREAK
import zipfile
import io
from datetime import datetime
import argparse
from typing import Any, Dict, Optional, List
from pathlib import Path
import copy
import tempfile
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION
from docxcompose.composer import Composer


# ================= CONFIG =================

REQUIRED_COLUMNS = [
    "Title",
    "Fname",
    "Sname",
    "Business Name",
    "Add1",
    "Add2",
    "Town",
    "County",
    "Post Code",
]

# ================= HELPERS =================

def safe_filename(name: str) -> str:
    """Create safe filename from business name."""
    if not name or not isinstance(name, str):
        return "letter"
    return "".join(c if c.isalnum() or c in " _-" else "_" for c in name).strip()

def replace_placeholders_properly(doc: Document, mapping: dict):
    """
    Replace placeholders while PRESERVING formatting.
    This is the key fix - we iterate through runs, not paragraphs.
    """
    for paragraph in doc.paragraphs:
        # Build full paragraph text to check for placeholders
        full_text = paragraph.text
        
        # Check if any placeholder exists in this paragraph
        has_placeholder = any(key in full_text for key in mapping.keys())
        
        if has_placeholder:
            # We need to replace text while preserving runs
            # Strategy: concatenate all run texts, do replacement, then redistribute
            
            # Collect all runs and their properties
            runs_data = []
            for run in paragraph.runs:
                runs_data.append({
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name,
                    'font_size': run.font.size,
                    'font_color': run.font.color.rgb if run.font.color.rgb else None
                })
            
            # Get full text and perform replacements
            full_text = ''.join([r['text'] for r in runs_data])
            for key, value in mapping.items():
                full_text = full_text.replace(key, value)
            
            # Clear all runs
            for _ in range(len(paragraph.runs)):
                paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)
            
            # Add new run with replaced text, using first run's formatting as template
            if runs_data:
                new_run = paragraph.add_run(full_text)
                template = runs_data[0]
                new_run.bold = template['bold']
                new_run.italic = template['italic']
                new_run.underline = template['underline']
                if template['font_name']:
                    new_run.font.name = template['font_name']
                if template['font_size']:
                    new_run.font.size = template['font_size']
                if template['font_color']:
                    new_run.font.color.rgb = template['font_color']
            else:
                paragraph.add_run(full_text)
    
    # Also handle tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = paragraph.text
                    has_placeholder = any(key in full_text for key in mapping.keys())
                    
                    if has_placeholder:
                        # Same strategy for table cells
                        runs_data = []
                        for run in paragraph.runs:
                            runs_data.append({
                                'text': run.text,
                                'bold': run.bold,
                                'italic': run.italic,
                                'underline': run.underline,
                                'font_name': run.font.name,
                                'font_size': run.font.size,
                                'font_color': run.font.color.rgb if run.font.color.rgb else None
                            })
                        
                        full_text = ''.join([r['text'] for r in runs_data])
                        for key, value in mapping.items():
                            full_text = full_text.replace(key, value)
                        
                        # Clear runs
                        for _ in range(len(paragraph.runs)):
                            paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)
                        
                        # Add new run
                        if runs_data:
                            new_run = paragraph.add_run(full_text)
                            template = runs_data[0]
                            new_run.bold = template['bold']
                            new_run.italic = template['italic']
                            new_run.underline = template['underline']
                            if template['font_name']:
                                new_run.font.name = template['font_name']
                            if template['font_size']:
                                new_run.font.size = template['font_size']
                            if template['font_color']:
                                new_run.font.color.rgb = template['font_color']
                        else:
                            paragraph.add_run(full_text)

def prepare_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    """
    Clean and validate DataFrame.
    Returns DataFrame with properly named columns.
    """
    # Clean column names
    df.columns = df.columns.str.strip()
    
    # Remove unnamed columns
    df = df.loc[:, ~df.columns.str.contains("^Unnamed")]
    
    # Fill NaN with empty string
    df = df.fillna("")
    
    # Case-insensitive header matching
    header_map = {col.lower(): col for col in df.columns}
    
    # Check for required columns
    missing = []
    for req_col in REQUIRED_COLUMNS:
        req_lower = req_col.lower()
        if req_lower not in header_map:
            missing.append(req_col)
    
    if missing:
        raise ValueError(f"Missing required columns: {missing}")
    
    # Create DataFrame with properly named columns
    clean_data = {}
    for req_col in REQUIRED_COLUMNS:
        req_lower = req_col.lower()
        source_col = header_map[req_lower]
        clean_data[req_col] = df[source_col].astype(str).str.strip()
    
    return pd.DataFrame(clean_data)


def remove_trailing_empty_paragraphs(doc):
    while doc.paragraphs and not doc.paragraphs[-1].text.strip():
        p = doc.paragraphs[-1]._element
        p.getparent().remove(p)


def combine_letters_from_individual(letter_docs: List[bytes]) -> bytes:
    if not letter_docs:
        return b""

    # Load first document as base
    master = Document(io.BytesIO(letter_docs[0]))
    composer = Composer(master)

    # Append remaining documents
    for doc_bytes in letter_docs[1:]:
        doc = Document(io.BytesIO(doc_bytes))
        composer.append(doc)

    buffer = io.BytesIO()
    master.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# def combine_letters(letters: List[bytes]) -> bytes:
    if not letters:
        return b""

    combined_doc = Document(io.BytesIO(letters[0]))

    # Clean trailing empty paragraphs from first letter
    remove_trailing_empty_paragraphs(combined_doc)

    for i in range(1, len(letters)):
        letter_doc = Document(io.BytesIO(letters[i]))

        # IMPORTANT: add SECTION break (not page break)
        combined_doc.add_section(WD_SECTION.NEW_PAGE)

        # Copy paragraphs
        for paragraph in letter_doc.paragraphs:
            if not paragraph.text.strip() and len(paragraph.runs) == 0:
                continue
            new_para = combined_doc.add_paragraph()
            if paragraph.style:
                new_para.style = paragraph.style
            for run in paragraph.runs:
                new_run = new_para.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                if run.font.name:
                    new_run.font.name = run.font.name
                if run.font.size:
                    new_run.font.size = run.font.size

        # Copy tables
        for table in letter_doc.tables:
            new_table = combined_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    new_cell = new_table.cell(r_idx, c_idx)
                    new_cell.text = cell.text

        remove_trailing_empty_paragraphs(combined_doc)

    buffer = io.BytesIO()
    combined_doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# ================= WEB SERVICE CLASS =================

class LetterGenerationService:
    def __init__(self, template_path: str = None):
        if template_path is None:
            # Try to find template in common locations
            possible_paths = [
                "app/templates/letters/letter_template.docx",
                "LetterTemplate.docx",
                Path(__file__).parent.parent / "templates" / "letters" / "letter_template.docx",
            ]
            for path in possible_paths:
                if isinstance(path, Path):
                    path_str = str(path)
                else:
                    path_str = path
                
                if os.path.exists(path_str):
                    template_path = path_str
                    print(f"✓ Found template at: {template_path}")
                    break
        
        if template_path is None or not os.path.exists(template_path):
            raise FileNotFoundError(
                "Template file not found. Please provide a template path or place "
                "LetterTemplate.docx in the working directory."
            )
        
        self.template_path = template_path
    
    def generate_from_excel(self, excel_path: str, mode: str = "zip", 
                          letters_per_file: int = 1) -> dict[str, Any]:
        """Generate letters from Excel file."""
        try:
            if not os.path.exists(excel_path):
                return {"error": f"Excel file not found: {excel_path}"}
            
            df = pd.read_excel(excel_path, dtype=str)
            return self.generate_from_dataframe(df, mode, letters_per_file)
        except Exception as e:
            return {"error": f"Cannot read Excel file: {str(e)}"}
    
    def generate_from_csv(self, csv_path: str, mode: str = "zip", 
                         letters_per_file: int = 1) -> dict[str, Any]:
        """Generate letters from CSV file."""
        try:
            if not os.path.exists(csv_path):
                return {"error": f"CSV file not found: {csv_path}"}
            
            df = pd.read_csv(csv_path, dtype=str)
            return self.generate_from_dataframe(df, mode, letters_per_file)
        except Exception as e:
            return {"error": f"Cannot read CSV file: {str(e)}"}
    
    def generate_from_file(self, file_path: str, mode: str = "zip", 
                          letters_per_file: int = 1) -> dict[str, Any]:
        """Generate letters from Excel or CSV file."""
        try:
            if file_path.lower().endswith(('.xlsx', '.xls')):
                return self.generate_from_excel(file_path, mode, letters_per_file)
            elif file_path.lower().endswith('.csv'):
                return self.generate_from_csv(file_path, mode, letters_per_file)
            else:
                return {"error": "Unsupported file format. Use Excel (.xlsx, .xls) or CSV (.csv)"}
        except Exception as e:
            return {"error": f"Cannot read file: {str(e)}"}
    
    def generate_from_dataframe(self, df: pd.DataFrame, mode: str = "zip", 
                               letters_per_file: int = 1) -> dict[str, Any]:
        """
        Generate letters from DataFrame.
        
        Args:
            df: Source DataFrame
            mode: "zip" = ZIP with individual files (1 per DOCX), "combined" = ZIP with N letters per DOCX
            letters_per_file: How many letters per DOCX file (only used in "combined" mode)
        
        NOTE: ALL rows in the dataframe will be processed. No limit parameter.
        """
        try:
            df = prepare_dataframe(df)
        except ValueError as e:
            return {"error": str(e)}
        
        # Process ALL rows - no limit
        total_letters = len(df)
        
        if mode == "zip":
            # Generate ZIP with one letter per DOCX file
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for idx, row in df.iterrows():
                    data = row.to_dict()
                    
                    # Load fresh template for each letter
                    doc = Document(self.template_path)
                    
                    # Create replacements mapping
                    replacements = {f"{{{k}}}": v for k, v in data.items()}
                    replacements["{Title} {Sname}"] = f"{data.get('Title', '')} {data.get('Sname', '')}".strip()
                    
                    # Replace placeholders properly
                    replace_placeholders_properly(doc, replacements)
                    
                    # Save to bytes for ZIP
                    letter_buffer = io.BytesIO()
                    doc.save(letter_buffer)
                    letter_buffer.seek(0)
                    
                    filename = safe_filename(data.get("Business Name", "")) or f"letter_{idx+1}"
                    zipf.writestr(f"{filename}.docx", letter_buffer.getvalue())
            
            zip_buffer.seek(0)
            return {
                "content": zip_buffer.getvalue(),
                "filename": f"letters_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
                "content_type": "application/zip",
                "total_letters": total_letters,
                "files_created": total_letters
            }
        
        elif mode == "combined":
            # Generate ZIP with multiple DOCX files, N letters per file
            if letters_per_file < 1:
                return {"error": "letters_per_file must be at least 1"}
            
            print(f"DEBUG: Generating in combined mode with {letters_per_file} letters per file")
            print(f"DEBUG: Total letters: {total_letters}")
    
            # STEP 1: Generate ALL individual letters first (preserves template perfectly)
            all_letters: List[bytes] = []
    
            for i, (_, row) in enumerate(df.iterrows()):
                data = row.to_dict()
        
                # Generate single letter using template (same as individual mode)
                doc = Document(self.template_path)
                replacements = {f"{{{k}}}": v for k, v in data.items()}
                replacements["{Title} {Sname}"] = f"{data.get('Title', '')} {data.get('Sname', '')}".strip()
                replace_placeholders_properly(doc, replacements)
        
                # Save to bytes
                letter_buffer = io.BytesIO()
                doc.save(letter_buffer)
                letter_buffer.seek(0)
        
                all_letters.append(letter_buffer.getvalue())
    
            print(f"DEBUG: Generated {len(all_letters)} individual letters before batching")
    
            # STEP 2: Batch + combine using real DOCX merging
            from docxcompose.composer import Composer
    
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
                file_num = 1
    
                for start_idx in range(0, len(all_letters), letters_per_file):
                    batch = all_letters[start_idx:start_idx + letters_per_file]
    
                    print(f"DEBUG: Combining batch {file_num} with {len(batch)} letters "
                          f"(rows {start_idx+1} to {start_idx+len(batch)})")
    
                    # Load first document as master
                    master_doc = Document(io.BytesIO(batch[0]))
                    composer = Composer(master_doc)
    
                    # Append remaining documents in batch WITH forced section break
                    for letter_bytes in batch[1:]:
                        temp_doc = Document(io.BytesIO(letter_bytes))

                        # FORCE each new letter to start on a new page (real Word section break)
                        master_doc.add_section(WD_SECTION.NEW_PAGE)
                        composer.append(temp_doc)

                    # Save combined DOCX to bytes
                    combined_buffer = io.BytesIO()
                    master_doc.save(combined_buffer)
                    combined_buffer.seek(0)
    
                    zipf.writestr(f"letters_batch_{file_num}.docx", combined_buffer.getvalue())
    
                    file_num += 1
    
            zip_buffer.seek(0)
            num_files = file_num - 1
            
            print(f"DEBUG: Created {num_files} combined files")
    
            return {
                "content": zip_buffer.getvalue(),
                "filename": f"letters_batches_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
                "content_type": "application/zip",
                "total_letters": total_letters,
                "files_created": num_files,
                "letters_per_file": letters_per_file
            }
        
        else:
            return {"error": f"Unsupported mode: {mode}. Use 'zip' or 'combined'"}


# ================= MAIN (CLI) =================

def main():
    parser = argparse.ArgumentParser(description="Generate business acquisition letters from data file.")
    parser.add_argument("input_file", help="Path to input Excel or CSV file")
    parser.add_argument("--template", "-t", help="Path to template DOCX file", default="LetterTemplate.docx")
    parser.add_argument("--output", "-o", help="Output directory", default="Letters_Folder")
    parser.add_argument("--mode", "-m", choices=["individual", "zip", "combined"], default="individual",
                       help="individual/zip=separate files, combined=multiple letters per file")
    parser.add_argument("--letters-per-file", "-lpf", type=int, default=1,
                       help="Letters per DOCX file (only for 'combined' mode)")
    
    args = parser.parse_args()
    
    if not os.path.exists(args.template):
        sys.exit(f"Template file not found: {args.template}")
    
    if not os.path.exists(args.input_file):
        sys.exit(f"Input file not found: {args.input_file}")
    
    # Initialize service
    service = LetterGenerationService(args.template)
    
    # Generate based on mode
    if args.mode in ["individual", "zip"]:
        result = service.generate_from_file(args.input_file, "zip", 1)
    else:
        result = service.generate_from_file(args.input_file, "combined", args.letters_per_file)
    
    if "error" in result:
        sys.exit(f"Error: {result['error']}")
    
    # Save output
    os.makedirs(args.output, exist_ok=True)
    output_path = os.path.join(args.output, result['filename'])
    
    with open(output_path, 'wb') as f:
        f.write(result['content'])
    
    print(f"✓ Generated {result['total_letters']} letters in {result.get('files_created', 1)} file(s): {output_path}")

if __name__ == "__main__":
    main()